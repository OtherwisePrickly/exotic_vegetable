import os
from dotenv import load_dotenv
from flask import Flask, render_template, request, redirect, url_for, session, flash, send_from_directory, jsonify, send_file
from flask_session import Session
from flask_wtf.csrf import CSRFProtect
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired, Email
import boto3
import json
from docx import Document as DocxDocument
import io
from functools import wraps
import base64
import requests
from werkzeug.utils import secure_filename
from bs4 import BeautifulSoup
import hmac
import hashlib
import time

# Load environment variables from .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Use environment variables for sensitive data
app.secret_key = os.getenv('SECRET_KEY', 'default_secret_key')

# Initialize CSRF protection
csrf = CSRFProtect(app)

# Configure Flask-Session
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = os.path.join(app.instance_path, 'sessions')
app.config['SESSION_FILE_THRESHOLD'] = 100
app.config['SESSION_FILE_MODE'] = 0o600
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_USE_SIGNER'] = True
app.config['SESSION_KEY_PREFIX'] = 'your_app_session:'
Session(app)

# Ensure session directory exists
session_dir = app.config['SESSION_FILE_DIR']
os.makedirs(session_dir, exist_ok=True)

# Configure file upload folder
app.config['UPLOAD_FOLDER'] = os.path.join(app.instance_path, 'uploads')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize the Amazon Bedrock client with region from environment variables
client = boto3.client('bedrock-runtime', region_name=os.getenv('AWS_REGION', 'us-east-1'))

# Replace with your specific Bedrock model ID from environment variables
MODEL_ID = os.getenv('MODEL_ID')

# Amazon SES configuration for magic link
SES_CLIENT = boto3.client('ses', region_name=os.getenv('AWS_REGION'))
SENDER_EMAIL = os.getenv('SENDER_EMAIL')

# Set the allowed domain for magic link emails from environment variable
ALLOWED_DOMAINS = os.getenv('ALLOWED_DOMAINS')



def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Magic Link Functionality
def generate_magic_link(email):
    timestamp = int(time.time())
    secret_key = app.secret_key
    token = hmac.new(secret_key.encode(), f"{email}{timestamp}".encode(), hashlib.sha256).hexdigest()
    return url_for('magic_login', email=email, token=token, _external=True)

def send_magic_link(email):
    magic_link = generate_magic_link(email)
    response = SES_CLIENT.send_email(
        Source=SENDER_EMAIL,
        Destination={'ToAddresses': [email]},
        Message={
            'Subject': {'Data': 'Your Magic Login Link'},
            'Body': {
                'Text': {'Data': f"Click on the following link to log in: {magic_link}"}
            }
        }
    )
    return response

@app.route('/magic_login', methods=['GET'])
def magic_login():
    email = request.args.get('email')
    token = request.args.get('token')

    if not email or not token:
        flash('Invalid magic link. Please try again.', 'error')
        return redirect(url_for('login'))

    secret_key = app.secret_key
    timestamp = int(time.time())
    time_window = 600

    for t in range(timestamp - time_window, timestamp + 1):
        expected_token = hmac.new(secret_key.encode(), f"{email}{t}".encode(), hashlib.sha256).hexdigest()
        if hmac.compare_digest(expected_token, token):
            session['logged_in'] = True
            flash('Successfully logged in via magic link!', 'success')
            return redirect(url_for('index'))

    flash('The magic link has expired or is invalid. Please request a new one.', 'error')
    return redirect(url_for('login'))

@app.route('/request_magic_link', methods=['GET', 'POST'])
def request_magic_link():
    form = MagicLinkForm()
    if form.validate_on_submit():
        email = form.email.data
        domain = email.split('@')[-1]

        if domain not in ALLOWED_DOMAINS:
            flash(f"Only email addresses within the allowed domains are permitted.", 'error')
            return redirect(url_for('request_magic_link'))

        try:
            send_magic_link(email)
            flash('Magic link sent! Please check your email.', 'success')
        except Exception as e:
            flash(f'Error sending email: {e}', 'error')

    return render_template('request_magic_link.html', form=form)

def validate_with_bedrock(attributes=None, image_base64=None):
    content = ""

    if attributes:
        content += f"Is the following text relevant to generating a product description for clothing? Answer with 'yes' or 'no'. Text: {attributes} "

    if image_base64:
        content += "Is the content of this image related to clothing? Answer with 'yes' or 'no'."

    if not content:
        return False

    messages = [{
        "role": "user",
        "content": [
            {
                "type": "text",
                "text": content
            }
        ]
    }]

    if image_base64:
        messages[0]["content"].append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": "image/jpeg",
                "data": image_base64,
            }
        })

    data = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 10,
        "messages": messages
    }

    try:
        response = client.invoke_model(
            modelId=MODEL_ID,
            contentType='application/json',
            accept='application/json',
            body=json.dumps(data)
        )

        response_body = response['body'].read().decode('utf-8')
        result = json.loads(response_body)

        for message in result.get('content', []):
            answer = message.get('text', '').strip().lower()
            if "yes" in answer:
                return True

        return False

    except Exception as e:
        print(f"An error occurred during validation: {e}")
        return False

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'docx', 'txt'}

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        username = form.username.data
        password = form.password.data

        if username == os.getenv('ADMIN_USERNAME', 'admin') and password == os.getenv('ADMIN_PASSWORD', 'your_password'):
            session['logged_in'] = True
            return redirect(url_for('index'))
        else:
            flash('Invalid credentials. Please try again.', 'error')
    
    return render_template('login.html', form=form)

@app.route('/downloads/<filename>')
def download_file(filename):
    directory = os.path.join(app.static_folder, 'downloads')
    return send_from_directory(directory=directory, path=filename, as_attachment=True)

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    return redirect(url_for('login'))

@app.route('/', methods=['GET'])
@login_required
def index():
    # Load environment variables
    logo_url = os.getenv('LOGO_URL')
    welcome_message = os.getenv('WELCOME_MESSAGE')

    # Load session data
    products = session.get('products', [])
    bedrock_response = session.get('bedrock_response', None)
    translations = session.get('translations', {})
    selected_languages = session.get('selected_languages', [])

    return render_template('index.html', logo_url=logo_url, welcome_message=welcome_message, products=products, bedrock_response=bedrock_response, translations=translations, selected_languages=selected_languages)

@app.route('/upload', methods=['POST'])
@login_required
def upload():
    if 'file' not in request.files:
        flash('No file part', 'error')
        return redirect(request.url)

    file = request.files['file']
    if file.filename == '':
        flash('No selected file', 'error')
        return redirect(request.url)

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        try:
            file.save(file_path)

            if filename.endswith('.docx'):
                products = process_docx(file_path)
            elif filename.endswith('.txt'):
                products = process_txt(file_path)
            else:
                flash('Unsupported file type. Please upload a .docx or .txt file.', 'error')
                return redirect(url_for('index'))

            if not products:
                flash('The file must contain at least one product attribute or a valid image URL.', 'error')
                return redirect(url_for('index'))
            
            session['products'] = products
            session['bedrock_response'] = None
            session['translations'] = {}
            session.pop('selected_languages', None)
            flash('File successfully uploaded', 'success')

        except Exception as e:
            flash(f'Error saving file: {e}', 'error')
            return redirect(url_for('index'))
    else:
        flash('Unsupported file type. Please upload a .docx or .txt file.', 'error')

    return redirect(url_for('index'))

@app.route('/generate', methods=['POST'])
@login_required
def generate():
    try:
        products = session.get('products', [])

        welcome_message = os.getenv('WELCOME_MESSAGE')
        logo_url = os.getenv('LOGO_URL')

        if not products:
            flash("No products found. Please upload a document with product details or an image.", 'error')
            return redirect(url_for('index'))

        edited_products = []
        for index, product in enumerate(products):
            edited_key = f'edited_attributes_{index}'
            edited_content = request.form.get(edited_key)
            if edited_content:
                soup = BeautifulSoup(edited_content, 'html.parser')
                edited_attributes = [li.text for li in soup.find_all('li')]
                product['attributes'] = edited_attributes

            edited_products.append(product)
        
        session['products'] = edited_products

        for product in edited_products:
            if not product.get('attributes') and not product.get('image_base64'):
                flash("Product attributes or image is missing. Please ensure the document includes product details or a valid image.", 'error')
                return redirect(url_for('index'))

            attributes_content = ", ".join(product.get('attributes', []))
            is_valid = validate_with_bedrock(attributes=attributes_content, image_base64=product.get('image_base64'))

            if not is_valid:
                flash("The content does not appear to be related to clothing. Please check and try again.", 'error')
                return render_template('index.html', welcome_message=welcome_message, logo_url=logo_url, redirect_to_upload=True)

        bedrock_response = send_to_bedrock_model(edited_products)
        session['bedrock_response'] = bedrock_response

        languages = request.form.getlist('languages')
        translations = {}
        for language in languages:
            translations[language] = generate_translation(bedrock_response, language)

        session['translations'] = translations
        session['selected_languages'] = languages

    except ValueError as ve:
        flash(str(ve), 'error')
        return redirect(url_for('index'))
    except Exception as e:
        flash(f"An unexpected error occurred: {e}", 'error')
        return redirect(url_for('index'))

    return redirect(url_for('index'))

@app.route('/update_translation', methods=['POST'])
@login_required
def update_translation():
    data = request.get_json()
    language = data.get('language')
    add_translation = data.get('add')

    bedrock_response = session.get('bedrock_response', "")
    translations = session.get('translations', {})

    if add_translation:
        translations[language] = generate_translation(bedrock_response, language)
    else:
        if language in translations:
            del translations[language]

    session['translations'] = translations
    session['selected_languages'] = list(translations.keys())

    return jsonify(success=True)

@app.route('/download_descriptions', methods=['GET'])
@login_required
def download_descriptions():
    bedrock_response = session.get('bedrock_response')
    translations = session.get('translations', {})

    if not bedrock_response:
        flash('No product description available to download.', 'error')
        return redirect(url_for('index'))

    doc = DocxDocument()
    doc.add_heading('Product Description', level=1)
    doc.add_paragraph(bedrock_response)

    if translations:
        doc.add_heading('Translations', level=1)
        for language, translation in translations.items():
            doc.add_heading(f'{language} Translation', level=2)
            doc.add_paragraph(translation)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Product_Descriptions.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/reset', methods=['POST'])
@login_required
def reset():
    session.pop('products', None)
    session.pop('bedrock_response', None)
    session.pop('translations', None)
    session.pop('selected_languages', None)
    return redirect(url_for('index'))

def process_docx(file_path):
    try:
        document = DocxDocument(file_path)
        products = []
        paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        
        if not paragraphs:
            return products
        
        description = paragraphs[:-1] if paragraphs[:-1] else []
        image_url = paragraphs[-1] if paragraphs[-1].startswith('http') else None

        if not image_url and not description:
            description = [paragraphs[-1]]
        
        image_base64 = convert_image_to_base64(image_url) if image_url else None
        
        if description or image_base64:
            products.append({
                'attributes': description if description else [],
                'image_url': image_url,
                'image_base64': image_base64
            })
        
        return products
    except Exception as e:
        return []

def process_txt(file_path):
    products = []
    try:
        with open(file_path, 'r') as f:
            lines = [line.strip() for line in f if line.strip()]
            if not lines:
                return products
            
            description = lines[:-1] if lines[:-1] else []
            image_url = lines[-1] if lines[-1].startswith('http') else None
            
            if not image_url and not description:
                description = [lines[-1]]
            
            image_base64 = convert_image_to_base64(image_url) if image_url else None
            
            if description or image_base64:
                products.append({
                    'attributes': description if description else [],
                    'image_url': image_url,
                    'image_base64': image_base64
                })
    except Exception as e:
        pass
    
    return products

def convert_image_to_base64(image_url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"
        }
        response = requests.get(image_url, headers=headers)
        response.raise_for_status()

        return base64.b64encode(response.content).decode('utf-8')
    except requests.exceptions.RequestException as e:
        return None

import uuid

# Get description prompt and tone example from .env

def send_to_bedrock_model(products):
    messages = []

    # Load description_prompt and tone_example from the environment variables
    prompt_rules = os.getenv('PROMPT_RULES', 'Default prompt rules')
    description_prompt = os.getenv('DESCRIPTION_PROMPT', 'Default description prompt text')
    tone_example = os.getenv('TONE_EXAMPLE', 'Default tone example text')

    for product in products:
        unique_id = str(uuid.uuid4())

        combined_content = (
	    f"{prompt_rules} " # Add general prompt rules from the .env at the beginning
            f"{description_prompt} "  # Use the product description prompt from the .env
            f"Here is an example of the tone: {tone_example} "  # Use the tone example from the .env
        )

        if product.get('attributes'):
            combined_content += f"Attributes: {', '.join(product['attributes'])}. "
        if product.get('image_base64'):
            combined_content += "Use the provided image to enhance the product description. "

        combined_content += f"Unique ID: {unique_id}. "

        message = {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": combined_content
                }
            ]
        }

        if product.get('image_base64'):
            message["content"].append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/jpeg",
                    "data": product['image_base64'],
                }
            })

        messages.append(message)

    if not messages:
        raise ValueError("No valid content found to send to the Bedrock model.")

    data = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 1000,
        "messages": messages
    }

    try:
        response = client.invoke_model(
            modelId=MODEL_ID,
            contentType='application/json',
            accept='application/json',
            body=json.dumps(data)
        )

        response_body = response['body'].read().decode('utf-8')
        result = json.loads(response_body)

        content_array = result.get('content', [])
        summaries = "\n".join([item['text'] for item in content_array if item['type'] == 'text'])

        return summaries if summaries else "No result returned from Bedrock."
    except Exception as e:
        return f"An error occurred: {e}"

def generate_translation(description, language):
    prompt_text = (
        f"Translate the following English product description into {language}, ensuring that the translation is culturally and linguistically appropriate for a {language}-speaking audience. "
        f"The tone should be stylish, street-ready, and emphasize both quality and functionality, appealing to urban explorers. Please follow these guidelines to ensure the translation is both accurate and natural, minimizing the need for manual adjustments:\n\n"
        
        f"1. **Dominance and Impact**: Use language that conveys dominance and impact, such as strong action verbs like 'Dominate' or 'Conquer.' Adapt these terms to culturally appropriate equivalents in {language}.\n"
        
        f"2. **Tone and Word Choice**: Use modern, fashionable terms that reflect the product’s premium quality. Ensure the translation is engaging and direct. Avoid overly casual or regional terms unless they are widely accepted and enhance the product's appeal. Use consistent terminology throughout the description.\n"
        
        f"3. **Cultural Relevance and Specificity**: Ensure that the translation feels natural and relatable across different {language}-speaking regions. Replace generic phrases with specific, trendy alternatives that better capture the product’s stylish and urban appeal. For example, use 'Conquista' instead of 'Domina' in Spanish for stronger impact.\n"
        
        f"4. **Clarity and Simplicity**: Opt for clear, straightforward terms that maintain the stylish tone of the original. Simplify phrases where necessary to improve readability and accessibility, such as using 'Fabriquées' instead of 'Confectionnées' in French for better flow.\n"
        
        f"5. **Avoid Redundancy and Unnecessary Introductory Phrases**: Do not include phrases such as 'Voici la traduction' or 'Here is the translation.' The translation should start directly with the product description and avoid unnecessary words that don’t add value to the overall message.\n"
        
        f"6. **Final Output**: Provide only the final translated description, ensuring it is polished, culturally appropriate, and ready for direct use. There should be no additional explanations, notes, or suggestions included in the output.\n\n"
        
        f"Here is the product description:\n\n{description}"
    )

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": prompt_text
                }
            ]
        }
    ]

    data = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 1000,
        "messages": messages
    }

    try:
        response = client.invoke_model(
            modelId=MODEL_ID,
            contentType='application/json',
            accept='application/json',
            body=json.dumps(data)
        )

        response_body = response['body'].read().decode('utf-8')
        result = json.loads(response_body)

        content_array = result.get('content', [])
        translation = "\n".join([item['text'] for item in content_array if item['type'] == 'text'])

        if "A notice:" in translation or "Note:" in translation:
            translation = translation.split("A notice:")[0].strip()
            translation = translation.split("Note:")[0].strip()

        return translation if translation else f"No translation available for {language}."

    except Exception as e:
        return f"An error occurred while translating to {language}: {e}"

class MagicLinkForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    submit = SubmitField('Send Magic Link')

class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Login')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
